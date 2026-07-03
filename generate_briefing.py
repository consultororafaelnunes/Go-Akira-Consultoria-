"""
generate_briefing.py — Gera os 2 briefings consolidados ao final da fase consultiva.

Fluxo para cada cliente:
  1. Recebe lista de transcrições (saída de fetch_drive_transcripts)
  2. Sumariza cada reunião com Claude Haiku 4.5 (barato, rápido)
  3. Consolida em 2 briefings com Claude Opus 4.8 (qualidade máxima):
       - Briefing Jurídico   → equipe de Instrumentos Jurídicos / COF
       - Briefing de Manuais → consultoras de Manuais Operacionais
  4. Gera .docx a partir do template GoAkira
  5. Faz upload ao Drive como Google Doc editável em /<cliente>/Briefings/
"""

import os
from datetime import date
from io import BytesIO
from pathlib import Path

import anthropic
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

from drive_folders import get_client_folder_id, find_or_create_briefings_folder
from consultants import get_juridico_email, get_manuais_email


HAIKU_MODEL   = "claude-haiku-4-5-20251001"
OPUS_MODEL    = "claude-opus-4-8"
TEMPLATE_PATH = Path(__file__).parent / "Templete Ata de Reunião.docx"
GDOC_MIME     = "application/vnd.google-apps.document"

NAVY = RGBColor(0x1E, 0x27, 0x61)
GRAY = RGBColor(0x64, 0x74, 0x8B)

SCOPES = ["https://www.googleapis.com/auth/drive"]


# ── Credenciais ───────────────────────────────────────────────────────────────

def _get_credentials() -> Credentials:
    creds = Credentials(
        token=None,
        refresh_token=os.environ["GOOGLE_REFRESH_TOKEN"],
        client_id=os.environ["GOOGLE_CLIENT_ID"],
        client_secret=os.environ["GOOGLE_CLIENT_SECRET"],
        token_uri="https://oauth2.googleapis.com/token",
        scopes=SCOPES,
    )
    creds.refresh(Request())
    return creds


# ── Prompts ───────────────────────────────────────────────────────────────────

_JURIDICO_SYSTEM = """\
Você é um especialista em franchising e instrumentos jurídicos de franquia.
Analise os registros das reuniões de consultoria de Business Plan e extraia
TODAS as informações relevantes para elaboração dos instrumentos jurídicos
(COF e Contrato de Franquia). Seja completo, preciso e use linguagem objetiva.\
"""

_JURIDICO_PROMPT = """\
Projeto: {cliente} | {total} reuniões | Período: {periodo}

REGISTROS DAS REUNIÕES:
{summaries_text}

---
Elabore o BRIEFING JURÍDICO completo estruturado assim:

## 1. IDENTIFICAÇÃO DO NEGÓCIO
Nome da marca, segmento, categoria de franquia, descrição do modelo.

## 2. MODELO DE FRANQUIA
Formato das unidades, tamanho padrão, territórios, raio de exclusividade, meta de expansão.

## 3. ESTRUTURA FINANCEIRA
Taxa de franquia (valor e condições), royalties (%, base, periodicidade), fundo de marketing,
investimento inicial estimado, payback estimado, demais encargos.

## 4. OBRIGAÇÕES DO FRANQUEADOR
Suporte, treinamento, fornecimento de insumos, visitas técnicas, sistemas e plataformas.

## 5. OBRIGAÇÕES DO FRANQUEADO
Padrões operacionais obrigatórios, fornecedores homologados, treinamentos, relatórios, restrições.

## 6. PRAZO E RENOVAÇÃO
Prazo do contrato, condições de renovação, direito de preferência.

## 7. RESCISÃO E SAÍDA
Condições de rescisão (ambas as partes), não-concorrência, destino do ponto e estoque.

## 8. PROPRIEDADE INTELECTUAL
Marcas (registradas ou em processo), segredos de negócio, restrições pós-encerramento.

## 9. PONTOS DE ATENÇÃO JURÍDICA
Questões sensíveis ou não resolvidas que precisam de decisão antes de redigir os instrumentos.
Particularidades que demandam cláusulas específicas.

## 10. HISTÓRICO DE DECISÕES RELEVANTES
Principais decisões tomadas durante a consultoria. Mudanças de posição do cliente ao longo das reuniões.

## 11. DADOS FALTANTES
Liste os dados que o consultor jurídico precisará solicitar ao cliente.\
"""

_MANUAIS_SYSTEM = """\
Você é um especialista em processos operacionais e manuais de franquia.
Analise os registros das reuniões de consultoria de Business Plan e extraia
TODAS as informações necessárias para elaboração dos Manuais Operacionais.
Seja completo, detalhado e prático — as consultoras usarão este briefing
como base para estruturar cada manual.\
"""

_MANUAIS_PROMPT = """\
Projeto: {cliente} | {total} reuniões | Período: {periodo}

REGISTROS DAS REUNIÕES:
{summaries_text}

---
Elabore o BRIEFING DE MANUAIS OPERACIONAIS completo estruturado assim:

## 1. VISÃO GERAL DO NEGÓCIO
Conceito, proposta de valor, posicionamento de mercado, público-alvo, diferenciais competitivos.

## 2. PRODUTOS E SERVIÇOS
Portfólio completo, produtos obrigatórios x opcionais, fichas técnicas (ingredientes, quantidades,
preparo), padrões de qualidade e apresentação, fornecedores por categoria.

## 3. OPERAÇÃO DA UNIDADE
Horário de funcionamento, fluxo de atendimento (passo a passo), abertura e fechamento (checklist),
fechamento de caixa, gestão de estoque.

## 4. PADRÕES DE ATENDIMENTO
Script de atendimento, tratamento de reclamações, uniformes, higiene e limpeza.

## 5. INFRAESTRUTURA E EQUIPAMENTOS
Layout padrão, equipamentos obrigatórios (com specs), mobiliário, identidade visual, sistema de PDV.

## 6. GESTÃO DA EQUIPE
Estrutura de cargos, funções e responsabilidades, recrutamento e seleção, treinamento inicial,
treinamento contínuo.

## 7. MARKETING LOCAL
Identidade visual (cores, fontes, logotipo), materiais de comunicação, diretrizes para redes sociais,
ações de marketing local permitidas, eventos e promoções sazonais.

## 8. INDICADORES DE DESEMPENHO (KPIs)
Métricas operacionais, metas esperadas, relatórios que o franqueado deve enviar.

## 9. GESTÃO DE QUALIDADE E AUDITORIA
Critérios de auditoria, plano de ação para não-conformidades, vigilância sanitária.

## 10. PROCESSOS ESPECÍFICOS DO SEGMENTO
Processos únicos ou diferenciais, sazonalidade, particularidades regionais.

## 11. ESTRUTURA RECOMENDADA DOS MANUAIS
Com base nas informações acima:
- Quais manuais produzir (ex: Manual de Operações, Manual de Produto, Manual de Marketing, etc.)
- Ordem de prioridade de elaboração
- Capítulos recomendados para cada manual

## 12. DADOS FALTANTES
Informações que as consultoras precisarão levantar com o cliente antes de iniciar os manuais.\
"""


# ── Sumarização individual (Haiku) ────────────────────────────────────────────

def _summarize_for_briefing(api: anthropic.Anthropic, t: dict) -> str:
    """Extrai pontos relevantes de uma transcrição para alimentar os briefings."""
    assunto = t.get("assunto") or t.get("subject", "")
    data    = t.get("data_reuniao") or t.get("date", "")
    fase    = t.get("fase", "BP")
    texto   = t["transcript"][:8000]  # Haiku tem 200k ctx mas limitamos por custo

    prompt = f"""\
Reunião: {assunto}
Data: {data} | Fase: {fase}

TRANSCRIÇÃO:
{texto}

Extraia de forma estruturada e concisa:
1. Decisões tomadas (preserve valores, percentuais e datas exatos)
2. Informações sobre o modelo de negócio
3. Processos operacionais discutidos
4. Questões financeiras
5. Pendências e pontos em aberto
6. Dados relevantes para instrumentos jurídicos ou manuais operacionais

Preserve números e valores exatos. Seja objetivo.\
"""
    resp = api.messages.create(
        model=HAIKU_MODEL,
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}],
    )
    return resp.content[0].text


# ── Geração do briefing (Opus) ────────────────────────────────────────────────

def _generate_briefing_text(
    api: anthropic.Anthropic,
    cliente: str,
    summaries_text: str,
    total: int,
    periodo: str,
    tipo: str,
) -> str:
    system = _JURIDICO_SYSTEM if tipo == "juridico" else _MANUAIS_SYSTEM
    prompt = (_JURIDICO_PROMPT if tipo == "juridico" else _MANUAIS_PROMPT).format(
        cliente=cliente,
        total=total,
        periodo=periodo,
        summaries_text=summaries_text,
    )

    tipo_label = "Jurídico" if tipo == "juridico" else "de Manuais"
    print(f"      Gerando Briefing {tipo_label} com {OPUS_MODEL}...")

    resp = api.messages.create(
        model=OPUS_MODEL,
        max_tokens=8000,
        system=system,
        messages=[{"role": "user", "content": prompt}],
    )
    for block in resp.content:
        if block.type == "text":
            return block.text
    return ""


# ── Montagem do .docx ─────────────────────────────────────────────────────────

def _build_docx(
    cliente: str,
    tipo: str,
    text: str,
    periodo: str,
    total: int,
) -> bytes:
    """Gera o .docx do briefing usando o template GoAkira como base."""
    doc = Document(TEMPLATE_PATH)

    # Limpa o corpo preservando sectPr (margens, tamanho da página)
    body = doc.element.body
    WNS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    sect_pr = body.find(f"{WNS}sectPr")
    for child in list(body):
        if child is not sect_pr:
            body.remove(child)

    tipo_label = "Briefing Jurídico" if tipo == "juridico" else "Briefing de Manuais Operacionais"
    destino    = "Instrumentos Jurídicos" if tipo == "juridico" else "Manuais Operacionais"

    def _run_style(run, size: int, bold: bool = False, color: RGBColor = None):
        run.font.size = Pt(size)
        run.bold = bold
        if color:
            run.font.color.rgb = color

    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{tipo_label}\n{cliente}")
    _run_style(r, 16, bold=True, color=NAVY)
    p.paragraph_format.space_after = Pt(4)

    # Metadados
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Período: {periodo}  ·  {total} reuniões\n"
        f"Destinado a: Equipe de {destino}\n"
        f"Gerado em: {date.today().strftime('%d/%m/%Y')}"
    )
    _run_style(r, 9, color=GRAY)
    p.paragraph_format.space_after = Pt(10)

    # Corpo — parse simples de Markdown
    for line in text.split("\n"):
        stripped = line.strip()

        if not stripped:
            doc.add_paragraph().paragraph_format.space_after = Pt(2)
            continue

        if stripped.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[3:])
            _run_style(r, 12, bold=True, color=NAVY)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after  = Pt(3)

        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            r = p.add_run(stripped[4:])
            _run_style(r, 11, bold=True, color=NAVY)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after  = Pt(2)

        elif stripped.startswith(("- ", "• ")):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(12)
            r = p.add_run(f"• {stripped[2:]}")
            _run_style(r, 10)
            p.paragraph_format.space_after = Pt(1)

        else:
            p = doc.add_paragraph()
            r = p.add_run(stripped)
            _run_style(r, 10)
            p.paragraph_format.space_after = Pt(2)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _upload_gdoc(service, docx_bytes: bytes, filename: str, folder_id: str) -> dict:
    meta = {
        "name":     filename,
        "mimeType": GDOC_MIME,
        "parents":  [folder_id],
    }
    media = MediaIoBaseUpload(
        BytesIO(docx_bytes),
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        resumable=False,
    )
    return service.files().create(
        body=meta, media_body=media, fields="id, webViewLink",
        supportsAllDrives=True,
    ).execute()


# ── Ponto de entrada ──────────────────────────────────────────────────────────

def generate_briefings(
    cliente: str,
    transcripts: list[dict],
    root_folder_id: str | None = None,
) -> dict:
    """
    Gera os 2 briefings para um cliente e faz upload ao Drive.

    Args:
        cliente:         Nome canônico do cliente.
        transcripts:     Lista de dicts (saída de fetch_drive_transcripts).
        root_folder_id:  ID da pasta raiz no Drive (None = My Drive).

    Returns:
        {"juridico": {"file_id", "link"}, "manuais": {"file_id", "link"}}
    """
    if not transcripts:
        raise ValueError(f"Nenhuma transcrição para '{cliente}'")

    api = anthropic.Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])

    # Período
    dts = [t["data_dt"] for t in transcripts if hasattr(t.get("data_dt"), "strftime")]
    periodo = (
        f"{min(dts).strftime('%d/%m/%Y')} a {max(dts).strftime('%d/%m/%Y')}"
        if dts else "período não determinado"
    )
    total = len(transcripts)
    print(f"\n   {cliente}: {total} reunioes | {periodo}")

    # 1. Sumarizar com Haiku
    print(f"   [1/3] Sumarizando {total} reunioes com Haiku...")
    parts = []
    for i, t in enumerate(transcripts, 1):
        label = t.get("assunto") or t.get("subject", f"Reuniao {i}")
        data  = t.get("data_reuniao") or t.get("date", "")
        print(f"         [{i}/{total}] {label[:55]} ({data})")
        summary = _summarize_for_briefing(api, t)
        parts.append(f"### Reuniao {i}: {label} — {data}\n{summary}")
    summaries_text = "\n\n---\n\n".join(parts)

    # 2. Gerar cada briefing com Opus + upload
    creds   = _get_credentials()
    service = build("drive", "v3", credentials=creds)

    # Usa a pasta real do cliente no Shared Drive GoAkira
    cl_fld = get_client_folder_id(cliente)
    if not cl_fld:
        from drive_folders import PASTA_RAIZ_CLIENTES
        cl_fld = root_folder_id or PASTA_RAIZ_CLIENTES
        print(f"   Aviso: cliente '{cliente}' nao mapeado — usando raiz como fallback")

    br_fld = find_or_create_briefings_folder(service, cl_fld)

    # Resolve e-mails dos destinatários antes de gerar (Manuais pode pedir input)
    juridico_email = get_juridico_email()
    manuais_email  = get_manuais_email(cliente)  # pergunta no terminal se não mapeado

    results = {}
    for tipo in ("juridico", "manuais"):
        print(f"\n   [2/3] {tipo}...")
        text = _generate_briefing_text(api, cliente, summaries_text, total, periodo, tipo)

        print(f"   [3/3] Montando .docx e subindo ao Drive...")
        docx_bytes = _build_docx(cliente, tipo, text, periodo, total)

        # Salva cópia local
        safe = cliente.replace(" ", "_").replace("/", "-")
        local = Path(f"briefing_{safe}_{tipo}.docx")
        local.write_bytes(docx_bytes)
        print(f"         Local: {local}")

        tipo_label = "Juridico" if tipo == "juridico" else "Manuais"
        filename   = f"Briefing {tipo_label} — {cliente}"
        uploaded   = _upload_gdoc(service, docx_bytes, filename, br_fld)

        link = uploaded.get("webViewLink", "—")
        dest = juridico_email if tipo == "juridico" else manuais_email
        results[tipo] = {"file_id": uploaded["id"], "link": link, "destinatario": dest}
        print(f"         Drive: {link}")
        if dest:
            print(f"         Destinatário: {dest}")
        else:
            print(f"         Destinatário: não definido (e-mail não será enviado)")

    return results
