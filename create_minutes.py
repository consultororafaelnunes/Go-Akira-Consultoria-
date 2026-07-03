"""
create_minutes.py — Gera a ata de reunião em .docx localmente (python-docx),
a partir do template "Templete Ata de Reunião.docx", e salva na pasta correta
do cliente dentro do Shared Drive GoAkira.

Fluxo:
  1. Abre o template .docx como base (herda página A4, margens e fonte Montserrat)
  2. Limpa o corpo e preenche com os dados do resumo gerado pelo Claude
  3. Localiza a pasta do cliente em Clientes/1.Formatação e Outros (via drive_folders.py)
  4. Dentro do cliente, acha a subpasta Business Plan → Atas e Formalizações
  5. Sobe como Google Doc editável (Drive converte .docx na importação)
  6. Retorna o link do documento criado

Não usa a API do Google Docs — só o Drive (escopo drive).
Todas as chamadas ao Drive usam supportsAllDrives=True (Shared Drive GoAkira).
"""

import os
import tempfile
from datetime import date, datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from google.oauth2.credentials import Credentials
from google.auth.transport.requests import Request
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload

from drive_folders import resolve_atas_folder, get_client_folder_id, PASTA_RAIZ_CLIENTES
from consultants import get_bp_consultant, get_consultant_email
from send_email import send_ata_notification


# ── Configuração ──────────────────────────────────────────────────────────────

SCOPES = ["https://www.googleapis.com/auth/drive"]

# Template usado como base visual (página A4, margens, fonte Montserrat)
TEMPLATE_PATH = Path(__file__).parent / "Templete Ata de Reunião.docx"

# Identidade visual (mesma paleta do PDF e do PPTX)
FONT  = "Montserrat"
NAVY  = RGBColor(0x1E, 0x27, 0x61)
CORAL = RGBColor(0xDC, 0x26, 0x26)
BLACK = RGBColor(0x1E, 0x29, 0x3B)
GRAY  = RGBColor(0x64, 0x74, 0x8B)

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
GDOC_MIME = "application/vnd.google-apps.document"  # converte para Google Doc editável


def get_credentials() -> Credentials:
    creds = Credentials(
        token=None,
        refresh_token=os.environ["GOOGLE_REFRESH_TOKEN"],
        client_id=os.environ["GOOGLE_CLIENT_ID"],
        client_secret=os.environ["GOOGLE_CLIENT_SECRET"],
        token_uri="https://oauth2.googleapis.com/token",
        scopes=SCOPES,
    )
    creds.refresh(Request())
    return creds


# ── Helpers de Drive ──────────────────────────────────────────────────────────

def upload_docx_to_drive(
    drive_service, local_path: str, filename: str, folder_id: str, max_retries: int = 3
) -> tuple[str, str]:
    """
    Sobe o .docx e converte em Google Doc editável.
    Retorna (file_id, webViewLink). Faz até max_retries tentativas em caso de falha transitória.
    """
    import time
    last_exc: Exception | None = None
    for attempt in range(max_retries):
        try:
            media = MediaFileUpload(local_path, mimetype=DOCX_MIME, resumable=False)
            meta  = {"name": filename, "parents": [folder_id], "mimeType": GDOC_MIME}
            file  = drive_service.files().create(
                body=meta, media_body=media, fields="id,webViewLink",
                supportsAllDrives=True,
            ).execute()
            return file["id"], file.get("webViewLink", "")
        except Exception as exc:
            last_exc = exc
            if attempt < max_retries - 1:
                wait = 5 * (attempt + 1)
                print(f"   Aviso: falha no upload (tentativa {attempt + 1}/{max_retries}) — aguardando {wait}s...")
                time.sleep(wait)
    raise RuntimeError(f"Upload falhou após {max_retries} tentativas: {last_exc}")


# ── Geração do .docx ──────────────────────────────────────────────────────────

def _clear_body(doc: Document) -> None:
    """Remove parágrafos e tabelas do template, preservando a seção (A4/margens)."""
    body = doc.element.body
    for el in list(body):
        # mantém apenas o <w:sectPr> (propriedades de página)
        if el.tag.endswith("}sectPr"):
            continue
        body.remove(el)


def _style_run(run, size=10, bold=False, italic=False, color=BLACK):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def _add_paragraph(doc, text, size=10, bold=False, italic=False, color=BLACK,
                   align=None, space_before=0, space_after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    run = p.add_run(text)
    _style_run(run, size=size, bold=bold, italic=italic, color=color)
    return p


def _add_labeled(doc, label, content, color=BLACK, size=10, space_after=4):
    """Parágrafo no padrão do template: '**Rótulo:** conteúdo'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r1 = p.add_run(f"{label}: ")
    _style_run(r1, size=size, bold=True, color=NAVY)
    r2 = p.add_run(content)
    _style_run(r2, size=size, bold=False, color=color)
    return p


def _add_section_title(doc, text):
    _add_paragraph(doc, text, size=11, bold=True, color=NAVY,
                   space_before=10, space_after=4)


def _add_bullets(doc, items, color=BLACK, empty_text="—"):
    if not items:
        _add_paragraph(doc, empty_text, size=10, italic=True, color=GRAY, space_after=2)
        return
    for item in items:
        try:
            p = doc.add_paragraph(style="List Bullet")
        except KeyError:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            item = f"•  {item}"
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(item)
        _style_run(run, size=10, color=color)


def build_ata_docx(summary: dict, output_path: str, template_path: str | Path = TEMPLATE_PATH) -> str:
    """
    Gera a ata em .docx a partir do template e do resumo. Retorna output_path.
    """
    template_path = Path(template_path)
    if template_path.exists():
        doc = Document(str(template_path))
        _clear_body(doc)
    else:
        # Sem o template: cria documento em branco (formatação aplicada por run)
        print(f"   ⚠️  Template não encontrado em {template_path} — gerando layout padrão")
        doc = Document()

    cliente       = summary.get("cliente", "Cliente não identificado")
    titulo        = summary.get("titulo_reuniao", "Reunião")
    data          = summary.get("data_reuniao") or date.today().strftime("%d/%m/%Y")
    duracao       = summary.get("duracao_estimada", "—")
    participantes = summary.get("participantes", [])
    resumo        = summary.get("resumo", "—")
    acionaveis    = summary.get("acionaveis", [])
    proximos      = summary.get("proximos_passos", [])
    alertas       = summary.get("alertas", [])
    sentimento    = summary.get("sentimento", "neutro").capitalize()
    prioridade    = summary.get("prioridade", "media").capitalize()

    # Título centralizado (12pt, negrito) — como no template
    _add_paragraph(
        doc, f"Ata de Reunião: {titulo} — {cliente}",
        size=12, bold=True, color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10,
    )

    # Metadados
    parts_txt = ", ".join(participantes) if participantes else "—"
    _add_labeled(doc, "Data", data)
    _add_labeled(doc, "Duração", duracao)
    _add_labeled(doc, "Participantes", parts_txt)

    # Objetivo / Resumo executivo
    _add_section_title(doc, "Objetivo e Resumo Executivo")
    _add_paragraph(doc, resumo, size=10, space_after=6)

    # Principais Acionáveis
    _add_section_title(doc, "Principais Acionáveis")
    _add_bullets(doc, acionaveis, empty_text="Nenhuma ação registrada.")

    # Próximos Passos
    _add_section_title(doc, "Plano de Ação e Próximos Passos")
    _add_bullets(doc, proximos, empty_text="Nenhum próximo passo registrado.")

    # Alertas e Riscos (só aparece se houver)
    if alertas:
        _add_section_title(doc, "⚠ Alertas e Riscos")
        _add_bullets(doc, alertas, color=CORAL)

    # Rodapé
    _add_paragraph(
        doc,
        f"Sentimento: {sentimento}  ·  Prioridade: {prioridade}",
        size=9, color=GRAY, space_before=12, space_after=2,
    )
    _add_paragraph(
        doc,
        f"Documento gerado automaticamente em {datetime.now().strftime('%d/%m/%Y às %H:%M')} "
        "· Agente de Resumo de Reuniões",
        size=8, italic=True, color=GRAY, space_after=0,
    )

    doc.save(output_path)
    return output_path


# ── Função principal ───────────────────────────────────────────────────────────

def create_meeting_minutes(
    summary: dict,
    root_folder_id: str | None = None,
    template_path: str | Path = TEMPLATE_PATH,
) -> dict:
    """
    Gera a ata (.docx) e salva na pasta do cliente no Drive.

    Returns:
        dict com {file_id, link, folder_id, cliente, nome_arquivo}
    """
    cliente = summary.get("cliente", "Cliente não identificado")
    print(f"\n   Criando ata para: {cliente}")

    creds         = get_credentials()
    drive_service = build("drive", "v3", credentials=creds)

    data          = summary.get("data_reuniao") or date.today().strftime("%d/%m/%Y")
    assunto       = summary.get("assunto") or summary.get("titulo_reuniao", "Reuniao")
    data_filename = data.replace("/", "-")
    nome_arquivo  = f"Ata — {cliente} — {data_filename} — {assunto}"

    # 1. Resolve a pasta de destino:
    #    cliente → pasta no Drive → Business Plan → Atas e Formalizações
    atas_folder_id = resolve_atas_folder(drive_service, cliente)

    if not atas_folder_id:
        # Cliente não mapeado ou sem Business Plan — usa root_folder_id como fallback
        fallback = root_folder_id or PASTA_RAIZ_CLIENTES
        print(f"   Aviso: usando pasta fallback para '{cliente}' (nao mapeado em drive_folders.py)")
        atas_folder_id = fallback

    # 2. Gera o .docx em arquivo temporário
    fd, tmp_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        build_ata_docx(summary, tmp_path, template_path=template_path)
        print(f"   Ata gerada: {nome_arquivo}")

        # 3. Upload
        file_id, link = upload_docx_to_drive(drive_service, tmp_path, nome_arquivo, atas_folder_id)
        print(f"   Salva no Drive: {link}")

        # 4. Notifica o consultor responsável pelo BP do cliente
        consultor = get_bp_consultant(cliente)
        email     = get_consultant_email(consultor) if consultor else None
        if email:
            try:
                send_ata_notification(summary, link, email)
            except Exception as e:
                print(f"   Aviso: falha ao enviar notificação para {email}: {e}")
        else:
            print(f"   Aviso: consultor de BP não mapeado para '{cliente}' — notificação não enviada")
    finally:
        Path(tmp_path).unlink(missing_ok=True)

    return {
        "file_id":      file_id,
        "link":         link,
        "folder_id":    atas_folder_id,
        "cliente":      cliente,
        "nome_arquivo": nome_arquivo,
    }


def create_all_minutes(
    summaries: list[dict],
    root_folder_id: str | None = None,
    template_path: str | Path = TEMPLATE_PATH,
) -> list[dict]:
    """
    Cria atas para todas as reuniões do dia. Injeta 'ata_link' em cada summary.
    """
    if not summaries:
        return []

    root_id = root_folder_id or os.environ.get("DRIVE_ROOT_FOLDER_ID")

    print(f"\n📋 Criando {len(summaries)} ata(s) de reunião...")
    results = []

    for summary in summaries:
        try:
            result = create_meeting_minutes(summary, root_id, template_path)
            summary["ata_link"] = result["link"]
            results.append(result)
        except Exception as e:
            print(f"   ❌ Erro ao criar ata para {summary.get('cliente')}: {e}")
            summary["ata_link"] = None

    print(f"✅ {len(results)} ata(s) criada(s) com sucesso")
    return results
